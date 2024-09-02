#!/usr/bin/env python
# encoding: utf-8

import wx
import os
import json
from threading import Thread
from docx import Document
import traceback

class MyFrame(wx.Frame):
    def __init__(self, *args, **kw):
        super(MyFrame, self).__init__(*args, **kw)
        self.InitUI()

    def InitUI(self):
        panel = wx.Panel(self)
        vbox = wx.BoxSizer(wx.VERTICAL)

        self.label1 = wx.StaticText(panel, label="请选择 JSON 数据文件的位置：")
        vbox.Add(self.label1, flag=wx.ALL, border=10)

        self.dirPicker = wx.DirPickerCtrl(panel, path="/Users/bigyang/myapp/yiheyuan/json", message="点我选择")
        vbox.Add(self.dirPicker, flag=wx.EXPAND|wx.ALL, border=10)

        self.startBtn = wx.Button(panel, label="点击开始生成")
        self.startBtn.Bind(wx.EVT_BUTTON, self.OnStart)
        vbox.Add(self.startBtn, flag=wx.ALL|wx.CENTER, border=10)

        panel.SetSizer(vbox)
        self.Centre()

    def OnStart(self, event):
        json_dir = self.dirPicker.GetPath()
        if not os.path.exists(json_dir):
            wx.MessageBox("请选择有效的JSON文件夹路径", "错误", wx.OK | wx.ICON_ERROR)
            return

        json_files = [f for f in os.listdir(json_dir) if f.endswith('.json')]
        total_files = len(json_files)
        if total_files == 0:
            wx.MessageBox("所选文件夹中没有JSON文件", "错误", wx.OK | wx.ICON_ERROR)
            return

        self.output_dir = "/Users/bigyang/myapp/yiheyuan/ok/"
        os.makedirs(self.output_dir, exist_ok=True)

        # 创建并显示进度条对话框
        self.progress_dialog = wx.ProgressDialog(
            "文件处理进度",
            "正在处理 JSON 文件...",
            maximum=total_files,
            parent=self,
            style=wx.PD_AUTO_HIDE | wx.PD_ELAPSED_TIME
        )

        # 创建并启动线程来处理文件
        self.thread = Thread(target=self.ProcessFiles, args=(json_dir, json_files))
        self.thread.start()

    def ProcessFiles(self, json_dir, json_files):
        try:
            # 在后台线程中处理文件
            for i, json_file in enumerate(json_files):
                self.process_json(json_dir, json_file)
                # 更新进度
                wx.CallAfter(self.UpdateProgress, i + 1, len(json_files))

            # 处理完成后，显示完成对话框
            wx.CallAfter(self.OnFinish)
        except Exception as e:
            # 捕获所有异常并记录日志
            error_msg = f"处理文件时出错: {str(e)}\n{traceback.format_exc()}"
            print(error_msg)  # 可以替换为将错误写入日志文件
            wx.CallAfter(wx.MessageBox, error_msg, "错误", wx.OK | wx.ICON_ERROR)

    def process_json(self, json_dir, json_file):
        try:
            template_path = "/Users/bigyang/myapp/yiheyuan/word/temp.docx"
            doc = Document(template_path)

            json_path = os.path.join(json_dir, json_file)
            with open(json_path, 'r', encoding='utf-8') as f:
                data = json.load(f)

            placeholders = {
                "年": "year", "月": "month", "日": "day", "总登记号": "zongdengjihao",
                "分类号": "fenleihao", "名称": "mingcheng", "年代": "niandai", "件数": "jianshu",
                "单位": "danwei", "尺寸": "chicun", "重量": "zhongliang", "质地": "zhidi",
                "完残情况": "wancanqingkuang", "来源": "laiyuan", "入馆凭证号": "ruguanpingzhenghao",
                "注销凭证号": "zhuxiaopingzhenghao", "级别": "jibie", "备注": "beizhu"
            }

            for p in doc.paragraphs:
                for key, placeholder in placeholders.items():
                    if placeholder in p.text:
                        p.text = p.text.replace(placeholder, str(data.get(key, "")))

            # 更新表格中的占位符
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for key, placeholder in placeholders.items():
                            if placeholder in cell.text:
                                cell.text = cell.text.replace(placeholder, str(data.get(key, "")))

            output_file = os.path.join(self.output_dir, f"{os.path.splitext(json_file)[0]}.docx")
            doc.save(output_file)
        except Exception as e:
            raise RuntimeError(f"处理文件 {json_file} 时出错: {str(e)}") from e

    def UpdateProgress(self, current, total):
        # 更新进度条
        self.progress_dialog.Update(current, f"已处理 {current} / {total} 个文件")

    def OnFinish(self):
        # 关闭进度条对话框
        if self.progress_dialog:
            self.progress_dialog.Destroy()

        wx.MessageBox(f"生成结束！共计生成 {len(os.listdir(self.output_dir))} 个文件！", "提示", wx.OK | wx.ICON_INFORMATION)

class MyApp(wx.App):
    def OnInit(self):
        frame = MyFrame(None, title="JSON 生成 Word 文档工具", size=(600, 300))
        frame.Show(True)
        self.SetTopWindow(frame)
        return True

if __name__ == "__main__":
    app = MyApp()
    app.MainLoop()
