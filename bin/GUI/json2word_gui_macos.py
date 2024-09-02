#!/usr/bin/env python
# encoding: utf-8

import wx
import os
import json
from docx import Document
import traceback

class MyFrame(wx.Frame):
    def __init__(self, *args, **kw):
        super(MyFrame, self).__init__(*args, **kw)
        self.InitUI()

        # 绑定窗口关闭事件
        self.Bind(wx.EVT_CLOSE, self.OnClose)

    def InitUI(self):
        panel = wx.Panel(self)
        vbox = wx.BoxSizer(wx.VERTICAL)

        self.label1 = wx.StaticText(panel, label="请选择 JSON 数据文件的位置：")
        vbox.Add(self.label1, flag=wx.ALL, border=10)

        self.dirPicker = wx.DirPickerCtrl(panel, path="/Users/bigyang/myapp/yiheyuan/json", message="点我选择")
        vbox.Add(self.dirPicker, flag=wx.EXPAND|wx.ALL, border=10)

        self.startBtn = wx.Button(panel, label="点击开始生成")
        self.startBtn.Bind(wx.EVT_BUTTON, self.OnStart)
        vbox.Add(self.startBtn, flag=wx.ALL|wx.CENTER, border=10)

        panel.SetSizer(vbox)
        self.Centre()

    def OnStart(self, event):
        json_dir = self.dirPicker.GetPath()
        if not os.path.exists(json_dir):
            wx.MessageBox("请选择有效的JSON文件夹路径", "错误", wx.OK | wx.ICON_ERROR)
            return

        json_files = [f for f in os.listdir(json_dir) if f.endswith('.json')]
        total_files = len(json_files)
        if total_files == 0:
            wx.MessageBox("所选文件夹中没有JSON文件", "错误", wx.OK | wx.ICON_ERROR)
            return

        self.output_dir = "/Users/bigyang/myapp/yiheyuan/ok/"
        os.makedirs(self.output_dir, exist_ok=True)

        self.progress_dialog = wx.ProgressDialog(
            "文件处理进度",
            "正在处理 JSON 文件...",
            maximum=total_files,
            parent=self,
            style=wx.PD_AUTO_HIDE | wx.PD_ELAPSED_TIME
        )

        batch_size = 5  # 每次处理的文件数量
        for start in range(0, total_files, batch_size):
            end = min(start + batch_size, total_files)
            for json_file in json_files[start:end]:
                self.process_json(json_dir, json_file)
            self.UpdateProgress(end, total_files)

        self.OnFinish()

    def process_json(self, json_dir, json_file):
        try:
            template_path = "/Users/bigyang/myapp/yiheyuan/word/temp.docx"
            doc = Document(template_path)

            json_path = os.path.join(json_dir, json_file)
            with open(json_path, 'r', encoding='utf-8') as f:
                data = json.load(f)

            placeholders = {
                "年": "year", "月": "month", "日": "day", "总登记号": "zongdengjihao",
                "分类号": "fenleihao", "名称": "mingcheng", "年代": "niandai", "件数": "jianshu",
                "单位": "danwei", "尺寸": "chicun", "重量": "zhongliang", "质地": "zhidi",
                "完残情况": "wancanqingkuang", "来源": "laiyuan", "入馆凭证号": "ruguanpingzhenghao",
                "注销凭证号": "zhuxiaopingzhenghao", "级别": "jibie", "备注": "beizhu"
            }

            # 更新段落中的占位符
            for p in doc.paragraphs:
                for key, placeholder in placeholders.items():
                    if placeholder in p.text:
                        p.text = p.text.replace(placeholder, str(data.get(key, "")))

            # 更新表格中的占位符
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for key, placeholder in placeholders.items():
                            if placeholder in cell.text:
                                cell.text = cell.text.replace(placeholder, str(data.get(key, "")))

            output_file = os.path.join(self.output_dir, f"{os.path.splitext(json_file)[0]}.docx")
            doc.save(output_file)
        except Exception as e:
            error_msg = f"处理文件 {json_file} 时出错: {str(e)}\n{traceback.format_exc()}"
            with open("/Users/bigyang/myapp/yiheyuan/log/json2word-errors.log", "a") as log_file:
                log_file.write(error_msg + "\n")

    def UpdateProgress(self, current, total):
        if current % 5 == 0:  # 每处理 10 个文件更新一次进度条
            self.progress_dialog.Update(current, f"已处理 {current} / {total} 个文件")

    def OnFinish(self):
        if self.progress_dialog:
            self.progress_dialog.Destroy()
        wx.MessageBox(f"生成结束！共计生成 {len(os.listdir(self.output_dir))} 个文件！", "提示", wx.OK | wx.ICON_INFORMATION)

    def OnClose(self, event):
        # 关闭时的处理
        if self.progress_dialog:
            self.progress_dialog.Destroy()
        self.Destroy()

class MyApp(wx.App):
    def OnInit(self):
        frame = MyFrame(None, title="JSON 生成 Word 文档工具", size=(400, 300))
        frame.Show(True)
        self.SetTopWindow(frame)
        return True

if __name__ == "__main__":
    app = MyApp()
    app.MainLoop()

