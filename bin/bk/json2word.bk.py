#!/usr/bin/env python
# encoding: utf-8

import json
import logging
import os
from datetime import datetime
from docx import Document

# 创建 log 和 ok 文件夹
os.makedirs('log', exist_ok=True)
os.makedirs('ok', exist_ok=True)

# 获取当前时间，生成日志文件名称
log_filename = datetime.now().strftime("log/%Y-%m-%d_%H-%M-%S.log")

# 配置日志，将日志写入文件
logging.basicConfig(
    filename=log_filename,
    level=logging.DEBUG,
    format='%(asctime)s - %(levelname)s - %(message)s'
)

# 读取 JSON 文件
json_filename = 'data.json'
try:
    with open(json_filename, 'r', encoding='utf-8') as f:
        data = json.load(f)
    logging.info("成功读取 JSON 文件")
except Exception as e:
    logging.error(f"读取 JSON 文件时出错: {e}")
    raise

# 打开一个 Word 文档
try:
    doc = Document('template.docx')
    logging.info("成功打开 Word 文档")
except Exception as e:
    logging.error(f"打开 Word 文档时出错: {e}")
    raise

# 定义占位符和 JSON 数据字段的映射
placeholder_mapping = {
    "[NAME]": data.get("name"),
    "[AGE]": data.get("age"),
    "[ADDRESS]": data.get("address"),
    "[PHONE]": data.get("phone")
}

# 替换段落中的占位符
for paragraph in doc.paragraphs:
    logging.debug(f"正在处理段落: {paragraph.text}")
    for placeholder, value in placeholder_mapping.items():
        if placeholder in paragraph.text:
            logging.info(f"找到占位符 {placeholder}，将其替换为: {value}")
            paragraph.text = paragraph.text.replace(placeholder, str(value))
        else:
            logging.debug(f"占位符 {placeholder} 未在段落中找到")

# 替换表格中的占位符
for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                logging.debug(f"正在处理表格单元格: {paragraph.text}")
                for placeholder, value in placeholder_mapping.items():
                    if placeholder in paragraph.text:
                        logging.info(f"找到占位符 {placeholder}，将其替换为: {value}")
                        paragraph.text = paragraph.text.replace(placeholder, str(value))
                    else:
                        logging.debug(f"占位符 {placeholder} 未在表格单元格中找到")

# 保存修改后的文档到 ok 文件夹中，名称与 json 文件一致
output_filename = os.path.join('ok', os.path.splitext(os.path.basename(json_filename))[0] + '.docx')
try:
    doc.save(output_filename)
    logging.info(f"成功保存修改后的文档到 {output_filename}")
except Exception as e:
    logging.error(f"保存文档时出错: {e}")
    raise

# 提示用户程序已完成
print("程序运行完毕，请查看结果文件。")

