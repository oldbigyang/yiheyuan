#!/usr/bin/env python
# encoding: utf-8

import os
import json
import logging
from docx import Document
from concurrent.futures import ProcessPoolExecutor, as_completed
from datetime import datetime
from rich.progress import Progress, TextColumn, BarColumn, TimeRemainingColumn
from multiprocessing import cpu_count

# 日志设置
log_dir = '/home/bigyang/python_bigyang/yiheyuan/log'
if not os.path.exists(log_dir):
    os.makedirs(log_dir)
log_file = os.path.join(log_dir, f'{datetime.now().strftime("%Y-%m-%d_%H-%M-%S")}.log')
logging.basicConfig(filename=log_file, level=logging.INFO, format='%(asctime)s - %(message)s')

# word 模板文件路径
template_path = "/home/bigyang/python_bigyang/yiheyuan/word/temp.docx"
# JSON 文件夹路径
json_folder = "/home/bigyang/python_bigyang/yiheyuan/json"
# 输出文件夹路径
output_folder = "/home/bigyang/python_bigyang/yiheyuan/ok"
if not os.path.exists(output_folder):
    os.makedirs(output_folder)

# 定义占位符与 JSON 数据的映射关系
def map_json_to_placeholders(data):
    return {
        "year": data.get("年"),
        "month": data.get("月"),
        "day": data.get("日"),
        "zongdengjihao": data.get("总登记号"),
        "fenleihao": data.get("分类号"),
        "name": data.get("名称"),
        "niandai": data.get("年代"),
        "jianshu": data.get("件数"),
        "danwei": data.get("单位"),
        "chicun": data.get("尺寸"),
        "zhongliang": data.get("重量"),
        "zhidi": data.get("质地"),
        "wancanqingkuang": data.get("完残情况"),
        "laiyuan": data.get("来源"),
        "ruguanpingzhenghao": data.get("入馆凭证号"),
        "zhuxiaopingzhenghao": data.get("注销凭证号"),
        "jibie": data.get("级别"),
        "beizhu": data.get("备注"),
        "fuzeren": data.get("负责人"),
        "danganbianhao": data.get("档案编号"),
        "xingzhuangneirongmiaoshu": data.get("形状内容描述"),
        "dangqianbaocuntiaojian": data.get("当前保存条件"),
        "mingjitiba": data.get("铭记题跋")
    }

# 处理单个 JSON 文件
def process_single_file(json_file):
    try:
        # 读取 JSON 数据
        with open(json_file, 'r', encoding='utf-8') as f:
            data = json.load(f)

        # 读取 Word 模板
        doc = Document(template_path)
        placeholders = map_json_to_placeholders(data)

        # 替换占位符
        for p in doc.paragraphs:
            for key, value in placeholders.items():
                if value:
                    p.text = p.text.replace(f'{key}', str(value))

        # 处理表格中的占位符
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for key, value in placeholders.items():
                        if value:
                            cell.text = cell.text.replace(f'{key}', str(value))

        # 保存生成的 Word 文件
        output_file = os.path.join(output_folder, os.path.basename(json_file).replace('.json', '.docx'))
        doc.save(output_file)

        logging.info(f'成功生成文件: {output_file}')
        return True
    except Exception as e:
        logging.error(f"处理文件 {json_file} 时出错: {str(e)}")
        return False

# 批量处理函数
def batch_process_json_files(json_files, batch_num, total_batches):
    total_files = len(json_files)
    with ProcessPoolExecutor(max_workers=cpu_count()) as executor:
        futures = [executor.submit(process_single_file, json_file) for json_file in json_files]
        
        # 使用 rich 进度条显示
        with Progress(
            TextColumn("[bold blue]{task.description}"),
            BarColumn(),
            "[progress.percentage]{task.percentage:>3.1f}%",
            TimeRemainingColumn(),
        ) as progress:
            task = progress.add_task(f"第 {batch_num}/{total_batches} 批文件处理进度", total=total_files)

            for future in as_completed(futures):
                future.result()  # 阻塞，确保每个任务完成
                progress.update(task, advance=1)

# 获取所有 JSON 文件并按文件名排序
def get_sorted_json_files(json_folder):
    json_files = [os.path.join(json_folder, f) for f in os.listdir(json_folder) if f.endswith('.json')]
    return sorted(json_files, key=lambda x: os.path.basename(x))  # 按文件名排序

if __name__ == "__main__":
    json_files = get_sorted_json_files(json_folder)
    
    if not json_files:
        logging.error("没有找到 JSON 文件")
        print("错误: 没有找到任何 JSON 文件。")
    else:
        # 计算总批次数量
        batch_size = 500
        total_batches = (len(json_files) + batch_size - 1) // batch_size  # 总批数,向上取整

        # 分批处理
        for i in range(0, len(json_files), batch_size):
            batch_num = (i // batch_size) + 1  # 当前批次
            batch = json_files[i:i + batch_size]
            batch_process_json_files(batch, batch_num, total_batches)

        print(f"程序运行完毕，共处理 {len(json_files)} 个文件，请查看生成的 Word 文件。")

