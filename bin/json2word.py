#!/usr/bin/env python
# encoding: utf-8

import os
import json
from datetime import datetime
from docx import Document
from tqdm import tqdm  # To add a progress bar
import logging

# 设置日志文件路径和名称
log_folder = "/Users/bigyang/myapp/yiheyuan/log"
os.makedirs(log_folder, exist_ok=True)
log_filename = os.path.join(log_folder, f"json2word_{datetime.now().strftime('%Y%m%d_%H%M%S')}.log")
logging.basicConfig(filename=log_filename, level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')

# 设置模板路径和输出文件夹
template_path = "/Users/bigyang/myapp/yiheyuan/word/temp.docx"
output_folder = "/Users/bigyang/myapp/yiheyuan/ok"
os.makedirs(output_folder, exist_ok=True)

# 占位符与JSON字段的映射
def replace_placeholders(doc, data):
    mapping = {
        "year": data.get("年"),
        "month": data.get("月"),
        "day": data.get("日"),
        "zongdengjihao": data.get("总登记号"),
        "fenleihao": data.get("分类号"),
        "name": data.get("名称"),
        "niandai": data.get("年代"),
        "jianshu": data.get("件数"),
        "danwei": data.get("单位"),
        "chicun": data.get("尺寸"),
        "zhongliang": data.get("重量"),
        "zhidi": data.get("质地"),
        "wancanqingkuang": data.get("完残情况"),
        "laiyuan": data.get("来源"),
        "ruguanpingzhenghao": data.get("入馆凭证号"),
        "zhuxiaopingzhenghao": data.get("注销凭证号"),
        "jibie": data.get("级别"),
        "beizhu": data.get("备注"),
        "fuzeren": data.get("负责人"),
        "danganbianhao": data.get("档案编号"),
        "xingzhuangneirongmiaoshu": data.get("形状内容描述"),
        "dangqianbaocuntiaojian": data.get("当前保存条件"),
        "mingjitiba": data.get("铭记题跋")

    }

    for paragraph in doc.paragraphs:
        for key, value in mapping.items():
            if key in paragraph.text:
                logging.info(f"Replacing placeholder: {key} with {value}")
                paragraph.text = paragraph.text.replace(f"{key}", str(value) if value else "")

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for key, value in mapping.items():
                    if key in cell.text:
                        logging.info(f"Replacing placeholder in table: {key} with {value}")
                        cell.text = cell.text.replace(f"{key}", str(value) if value else "")

# 单线程处理每个JSON文件
def process_single_file(json_filename):
    logging.info(f"Processing file: {json_filename}")
    with open(json_filename, 'r', encoding='utf-8') as f:
        data = json.load(f)
    
    # 读取模板文件
    doc = Document(template_path)
    
    # 替换占位符
    replace_placeholders(doc, data)
    
    # 生成 Word 文件
    output_filename = os.path.join(output_folder, os.path.basename(json_filename).replace(".json", ".docx"))
    doc.save(output_filename)
    logging.info(f"Word document saved as: {output_filename}")

# 主函数
def main():
    json_folder = input("请输入 JSON 文件夹路径: ")
    if not os.path.isdir(json_folder):
        print("输入的文件夹路径不存在，请重新输入。")
        logging.error("输入的文件夹路径不存在。")
        return
    
    json_files = [os.path.join(json_folder, f) for f in os.listdir(json_folder) if f.endswith('.json')]
    
    if not json_files:
        print("JSON 文件夹中没有找到任何 JSON 文件。")
        logging.error("未找到 JSON 文件。")
        return

    total_files = len(json_files)
    print(f"正在处理 {total_files} 个文件，请稍候...")
    logging.info(f"开始处理 {total_files} 个 JSON 文件。")

    # 单线程顺序处理每个文件
    for i, json_filename in enumerate(tqdm(json_files, desc="处理进度")):
        try:
            process_single_file(json_filename)
        except Exception as e:
            logging.error(f"处理文件 {json_filename} 时出错: {e}")
            print(f"处理文件 {json_filename} 时出错: {e}")
    
    print(f"程序运行完毕，一共生成 {total_files} 个文件，请查看。")
    logging.info(f"程序运行完毕，生成 {total_files} 个文件。")

if __name__ == "__main__":
    main()

